
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
stand = "Трудоваяденежкамозольная"
varname = '2.docx'
doc = Document(varname)
lol = stand.encode('koi8-r')
print(lol)
ou_str = []
for i in lol:
    ou_str.append(bin(i)[2:])
ist = ""
for i in ou_str:
    for j in i:
        ist = ist + j
print(ist)
text = ""
for i in doc.paragraphs:
    text = text + i.text
doc = Document()
p = doc.add_paragraph()
for i in range(len(text)):
    txt = text[i]
    if i < len(ist):
        if ist[i] == "1":
            run = p.add_run(txt).font.highlight_color = WD_COLOR_INDEX.TURQUOISE
        else:run = p.add_run(txt)
    else:
        run = p.add_run(txt)
doc.save('d.docx')